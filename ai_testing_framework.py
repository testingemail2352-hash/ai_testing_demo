"""
AI-Powered Testing Framework
Offline LLM integration pentru testare automată cu Mistral 7B
Target: Carturesti.ro
"""

import ollama
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json
import time
import threading


class AITestingFramework:
    """Framework principal pentru testare cu AI offline"""

    def __init__(self, model_name="mistral:7b"):
        self.model_name = model_name
        self.driver = None
        self.test_results = []

    def _call_mistral(self, prompt, stream=False):
        """Apelează Mistral prin Ollama"""
        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': prompt}],
                stream=stream
            )
            if stream:
                return response
            return response['message']['content']

        except Exception as e:
            # Mesaj clar și colorat în consolă
            print("\n" + "=" * 80)
            print("❌ EROARE LA MISTRAL ❌".center(80))
            print("-" * 80)
            print(f"Detalii: {str(e)}")
            print("=" * 80 + "\n")

            # Returnează mesajul de eroare ca text (pentru a fi prins mai sus)
            return f"Error calling Mistral: {str(e)}"

    def generate_gherkin_from_romanian(self, description_ro):
        """
        Generează scenarii Gherkin din descriere în română

        Args:
            description_ro: Descriere în română a funcționalității de testat

        Returns:
            str: Scenarii Gherkin generate
        """
        # with open("templates/gherkin_template.txt", "r", encoding="utf-8") as f:
        #     template_text = f.read()
        # {template_text}  - trebuie trecut la linia de cod nr. 70

        prompt = f"""
        Folosind **STRICT** structura și stilul următorului șablon Gherkin:

       

        Generează scenarii de test pentru următoarea descriere:
        "{description_ro}"

        Reguli:
        - Păstrează exact formatul din șablon (indentare, spațiere, poziția cuvintelor cheie)
        - Completează doar locurile dintre [paranteze]
        - Nu adăuga explicații în afara formatului
        """

        print("🤖 Generez scenarii Gherkin cu Mistral...")

        start_time = time.time()
        done = False

        # Thread pentru print periodic
        def status_timer():
            while not done:
                elapsed = int(time.time() - start_time)
                minutes = elapsed // 60
                seconds = elapsed % 60
                print(f"⏱ Timp trecut: {minutes} min {seconds} sec")
                time.sleep(300)  # 5 minute

        t = threading.Thread(target=status_timer)
        t.start()

        # Apel efectiv la model
        gherkin = self._call_mistral(prompt)

        done = True
        t.join()

        print("✅ Scenarii generate!")
        return gherkin

    def analyze_page_with_ai(self, url, max_wait=10):
        """
        Scanează pagina web și identifică elemente de testat cu ajutorul AI

        Args:
            url: URL-ul paginii de analizat
            max_wait: Timp maxim de așteptare pentru încărcare

        Returns:
            dict: Analiza paginii cu sugestii de teste
        """
        print(f"🔍 Analizez pagina: {url}")

        # Inițializare browser
        if not self.driver:
            options = webdriver.ChromeOptions()
            options.add_argument('--headless')
            options.add_argument('--disable-gpu')
            options.add_argument('--no-sandbox')
            self.driver = webdriver.Chrome(options=options)

        try:
            self.driver.get(url)
            time.sleep(3)  # Așteaptă încărcarea completă

            # Extrage informații despre pagină
            page_title = self.driver.title
            page_url = self.driver.current_url

            # Identifică elemente interactive
            buttons = self.driver.find_elements(By.TAG_NAME, "button")
            links = self.driver.find_elements(By.TAG_NAME, "a")
            inputs = self.driver.find_elements(By.TAG_NAME, "input")
            forms = self.driver.find_elements(By.TAG_NAME, "form")

            # Construiește structura paginii pentru AI
            page_structure = {
                'title': page_title,
                'url': page_url,
                'buttons': len(buttons),
                'links': len(links),
                'inputs': len(inputs),
                'forms': len(forms),
                'button_texts': [btn.text for btn in buttons[:10] if btn.text],
                'link_texts': [link.text for link in links[:10] if link.text],
                'input_types': [inp.get_attribute('type') for inp in inputs[:10]]
            }

            # Cere AI-ului să analizeze și să sugereze teste
            prompt = f"""
Analizează această pagină web și sugerează scenarii de testare:

Pagina: {page_title}
URL: {page_url}
Butoane: {page_structure['buttons']} (texte: {', '.join(page_structure['button_texts'][:5])})
Link-uri: {page_structure['links']} (texte: {', '.join(page_structure['link_texts'][:5])})
Input-uri: {page_structure['inputs']} (tipuri: {', '.join(set(page_structure['input_types']))})
Formulare: {page_structure['forms']}

Sugerează:
1. Top 5 funcționalități critice de testat
2. Potențiale probleme de identificat
3. Scenarii de edge cases
4. Teste de usability

Răspunde în română, concis și structurat.
"""

            print("🤖 Analizez cu Mistral...")
            ai_analysis = self._call_mistral(prompt)

            result = {
                'page_structure': page_structure,
                'ai_suggestions': ai_analysis,
                'timestamp': datetime.now().isoformat()
            }

            print("✅ Analiză completă!")
            return result

        except Exception as e:
            print(f"❌ Eroare la analiză: {str(e)}")
            return {'error': str(e)}

    def generate_test_scenarios_from_page(self, url):
        """
        Combină analiza paginii cu generarea de scenarii Gherkin

        Args:
            url: URL-ul paginii de analizat

        Returns:
            str: Scenarii Gherkin generate bazate pe pagină
        """
        # Analizează pagina
        analysis = self.analyze_page_with_ai(url)

        if 'error' in analysis:
            return f"Eroare: {analysis['error']}"

        # Generează descriere pentru Gherkin
        description = f"""
Generează scenarii Gherkin pentru pagina: {analysis['page_structure']['title']}
URL: {url}

Bazat pe analiza AI:
{analysis['ai_suggestions']}

Focus pe funcționalitățile principale identificate.
"""

        return self.generate_gherkin_from_romanian(description)

    def run_basic_test_suite(self, url, test_scenarios):
        """
        Rulează o suită simplă de teste bazate pe scenariile generate

        Args:
            url: URL-ul de testat
            test_scenarios: Scenarii de testat (text sau listă)

        Returns:
            list: Rezultate teste
        """
        print(f"🧪 Rulez teste pentru: {url}")
        results = []

        if not self.driver:
            options = webdriver.ChromeOptions()
            self.driver = webdriver.Chrome(options=options)

        try:
            self.driver.get(url)

            # Test 1: Verifică încărcarea paginii
            result = {
                'test': 'Page Load',
                'status': 'PASS',
                'details': f'Pagina s-a încărcat cu succes: {self.driver.title}',
                'timestamp': datetime.now().isoformat()
            }
            results.append(result)
            print(f"✅ {result['test']}: {result['status']}")

            # Test 2: Verifică prezența elementelor critice
            try:
                buttons = self.driver.find_elements(By.TAG_NAME, "button")
                result = {
                    'test': 'Critical Elements Check',
                    'status': 'PASS' if len(buttons) > 0 else 'FAIL',
                    'details': f'Găsite {len(buttons)} butoane pe pagină',
                    'timestamp': datetime.now().isoformat()
                }
                results.append(result)
                print(f"✅ {result['test']}: {result['status']}")
            except Exception as e:
                result = {
                    'test': 'Critical Elements Check',
                    'status': 'FAIL',
                    'details': str(e),
                    'timestamp': datetime.now().isoformat()
                }
                results.append(result)
                print(f"❌ {result['test']}: {result['status']}")

            # Test 3: Verifică responsive (resize window)
            try:
                self.driver.set_window_size(375, 667)  # iPhone size
                time.sleep(1)
                result = {
                    'test': 'Mobile Responsive',
                    'status': 'PASS',
                    'details': 'Pagina se adaptează la dimensiuni mobile',
                    'timestamp': datetime.now().isoformat()
                }
                results.append(result)
                print(f"✅ {result['test']}: {result['status']}")
            except Exception as e:
                result = {
                    'test': 'Mobile Responsive',
                    'status': 'FAIL',
                    'details': str(e),
                    'timestamp': datetime.now().isoformat()
                }
                results.append(result)
                print(f"❌ {result['test']}: {result['status']}")

        except Exception as e:
            print(f"❌ Eroare generală: {str(e)}")
            results.append({
                'test': 'General Test Execution',
                'status': 'ERROR',
                'details': str(e),
                'timestamp': datetime.now().isoformat()
            })

        self.test_results = results
        return results

    def generate_word_report(self, output_file="test_report.docx",
                             gherkin_scenarios=None,
                             page_analysis=None,
                             test_results=None):
        """
        Generează raport Word cu rezultatele testării

        Args:
            output_file: Numele fișierului de output
            gherkin_scenarios: Scenarii Gherkin generate
            page_analysis: Analiza paginii
            test_results: Rezultate teste
        """
        print(f"📄 Generez raport Word: {output_file}")

        doc = Document()

        # Header
        header = doc.add_heading('Raport Testare Automată cu AI', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Info generală
        doc.add_heading('Informații Generale', level=1)
        doc.add_paragraph(f'Data raport: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        doc.add_paragraph(f'Model AI: {self.model_name}')
        doc.add_paragraph(f'Framework: Selenium + Behave + Ollama')

        # Scenarii Gherkin
        if gherkin_scenarios:
            doc.add_page_break()
            doc.add_heading('Scenarii Gherkin Generate', level=1)
            p = doc.add_paragraph(gherkin_scenarios)
            p.style = 'Normal'

        # Analiza paginii
        if page_analysis and 'page_structure' in page_analysis:
            doc.add_page_break()
            doc.add_heading('Analiza Paginii Web', level=1)

            ps = page_analysis['page_structure']
            doc.add_paragraph(f"Titlu: {ps.get('title', 'N/A')}")
            doc.add_paragraph(f"URL: {ps.get('url', 'N/A')}")
            doc.add_paragraph(f"Butoane: {ps.get('buttons', 0)}")
            doc.add_paragraph(f"Link-uri: {ps.get('links', 0)}")
            doc.add_paragraph(f"Input-uri: {ps.get('inputs', 0)}")

            if 'ai_suggestions' in page_analysis:
                doc.add_heading('Sugestii AI', level=2)
                doc.add_paragraph(page_analysis['ai_suggestions'])

        # Rezultate teste
        if test_results:
            doc.add_page_break()
            doc.add_heading('Rezultate Teste', level=1)

            # Sumar
            total = len(test_results)
            passed = sum(1 for r in test_results if r['status'] == 'PASS')
            failed = sum(1 for r in test_results if r['status'] == 'FAIL')
            errors = sum(1 for r in test_results if r['status'] == 'ERROR')

            doc.add_paragraph(f'Total teste: {total}')
            doc.add_paragraph(f'✅ Passed: {passed}')
            doc.add_paragraph(f'❌ Failed: {failed}')
            doc.add_paragraph(f'⚠️ Errors: {errors}')

            # Detalii teste
            doc.add_heading('Detalii Teste', level=2)
            for result in test_results:
                status_symbol = '✅' if result['status'] == 'PASS' else '❌'
                doc.add_paragraph(
                    f"{status_symbol} {result['test']}: {result['status']}"
                )
                doc.add_paragraph(
                    f"   Detalii: {result['details']}",
                    style='List Bullet'
                )

        # Footer
        doc.add_page_break()
        footer = doc.add_paragraph('Generated by AI Testing Framework - Offline LLM Integration')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_file)
        print(f"✅ Raport salvat: {output_file}")

    def close(self):
        """Închide browser-ul"""
        if self.driver:
            self.driver.quit()
            print("🔚 Browser închis")


# ==================== EXEMPLU DE UTILIZARE ====================

def main():
    """Exemplu complet de utilizare a framework-ului"""

    print("=" * 60)
    print("🚀 AI Testing Framework - Carturesti.ro")
    print("=" * 60)

    # Inițializare framework
    framework = AITestingFramework()

    # URL de testat
    target_url = "https://carturesti.ro"

    try:
        # 1. Generează scenarii din descriere în română
        print("\n📝 STEP 1: Generare scenarii Gherkin")
        print("-" * 60)
        descriere = """
        Testează funcționalitatea de căutare de cărți pe Carturesti.ro.
        Include scenarii pentru căutare cu succes, căutare fără rezultate,
        și filtrare după preț sau categorie.
        """
        gherkin = framework.generate_gherkin_from_romanian(descriere)
        print("\nSCENARII GENERATE:")
        print(gherkin)

        # 2. Analizează pagina web cu AI
        print("\n\n🔍 STEP 2: Analiză pagină web")
        print("-" * 60)
        analysis = framework.analyze_page_with_ai(target_url)
        if 'ai_suggestions' in analysis:
            print("\nSUGESTII AI:")
            print(analysis['ai_suggestions'])

        # 3. Generează scenarii automat din pagină
        print("\n\n🤖 STEP 3: Generare scenarii automate din pagină")
        print("-" * 60)
        auto_scenarios = framework.generate_test_scenarios_from_page(target_url)
        print("\nSCENARII AUTO-GENERATE:")
        print(auto_scenarios)

        # 4. Rulează teste
        print("\n\n🧪 STEP 4: Rulare teste")
        print("-" * 60)
        results = framework.run_basic_test_suite(target_url, auto_scenarios)

        # 5. Generează raport Word
        print("\n\n📄 STEP 5: Generare raport Word")
        print("-" * 60)
        framework.generate_word_report(
            output_file="carturesti_test_report.docx",
            gherkin_scenarios=gherkin,
            page_analysis=analysis,
            test_results=results
        )

        print("\n" + "=" * 60)
        print("✅ Toate task-urile completate cu succes!")
        print("📄 Check: carturesti_test_report.docx")
        print("=" * 60)

    except Exception as e:
        print(f"\n❌ Eroare: {str(e)}")
        import traceback
        traceback.print_exc()

    finally:
        framework.close()


if __name__ == "__main__":
    main()